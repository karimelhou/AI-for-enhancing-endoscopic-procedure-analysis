from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Model Evaluation', level=1)

# Add a description
doc.add_paragraph(
    "Model evaluation is a critical phase in the development of AI models, particularly in the field of medical imaging, where accuracy and reliability are paramount. This section outlines the evaluation metrics, benchmarking methods, and comparative analysis employed to assess the performance of the classification and segmentation models developed in this study."
)

# Add a subheading for Classification Models
doc.add_heading('Classification Models', level=2)

# Add the classification benchmarking table
classification_table = doc.add_table(rows=1, cols=9)
classification_table.style = 'Table Grid'

# Define the columns for classification table
classification_columns = [
    "Model", "Train Accuracy", "Validation Accuracy", "Test Accuracy",
    "Precision", "Recall", "F1-Score", "AUC", "Loss"
]

# Add the header row for classification table
classification_hdr_cells = classification_table.rows[0].cells
for i, col in enumerate(classification_columns):
    classification_hdr_cells[i].text = col

# Define the data for classification table
classification_data = [
    ["EfficientNetB2", "98.06%", "98.06%", "98.06%", "98.19%", "98.06%", "98.12%", "99.82%", "0.0574"],
    ["EfficientNetB4", "99.97%", "99.80%", "100%", "-", "-", "-", "-", "-"]
]

# Add the data rows for classification table
for row_data in classification_data:
    row_cells = classification_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

# Add a subheading for Segmentation Models
doc.add_heading('Segmentation Models', level=2)

# Add the segmentation benchmarking table
segmentation_table = doc.add_table(rows=1, cols=7)
segmentation_table.style = 'Table Grid'

# Define the columns for segmentation table
segmentation_columns = [
    "Model", "Dice Coefficient", "Jaccard Index", "Loss"
]

# Add the header row for segmentation table
segmentation_hdr_cells = segmentation_table.rows[0].cells
for i, col in enumerate(segmentation_columns):
    segmentation_hdr_cells[i].text = col

# Define the data for segmentation table
segmentation_data = [
    ["DeepLabv3+", "0.89", "0.80", "0.15"],
    ["SegNet_transformer", "0.85", "0.75", "0.20"]
]

# Add the data rows for segmentation table
for row_data in segmentation_data:
    row_cells = segmentation_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

# Save the document
doc.save('Model_Evaluation.docx')
